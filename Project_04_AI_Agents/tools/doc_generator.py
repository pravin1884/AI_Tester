from docx import Document
import os

def create_test_plan(content: str, output_path: str = None) -> str:
    """Generate a .docx file containing the test plan content."""
    try:
        doc = Document()
        doc.add_heading('Test Plan', 0)
        
        # Simple splitting by paragraphs for formatting
        for p in content.split('\n'):
            if p.strip():
                if p.startswith('# '):
                    doc.add_heading(p.replace('# ', '').strip(), level=1)
                elif p.startswith('## '):
                    doc.add_heading(p.replace('## ', '').strip(), level=2)
                elif p.startswith('### '):
                    doc.add_heading(p.replace('### ', '').strip(), level=3)
                else:
                    doc.add_paragraph(p)
                    
        if not output_path:
            output_path = os.path.join('.tmp', 'Generated_Test_Plan.docx')
            
        # Ensure .tmp exists
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        
        doc.save(output_path)
        return output_path
    except Exception as e:
        print(f"Error generating document: {e}")
        return ""
